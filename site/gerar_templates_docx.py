# /// script
# requires-python = ">=3.11"
# dependencies = ["python-docx"]
# ///
"""
Gera os templates editáveis (.docx) para download a partir do conteúdo do site.

Fonte única de orientação: espelha o conteúdo de
conteudo/recursos/ripd-simplificado.md e model-card-educacional.md,
mas em formato preenchível (campos em branco, tabelas vazias).

Uso:
    .venv/bin/python site/gerar_templates_docx.py
    (ou: uv run site/gerar_templates_docx.py)

Saída: site/assets/downloads/*.docx  → publicados no build (copytree de assets/).
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = Path(__file__).resolve().parent / "assets" / "downloads"
ROXO = RGBColor(0x7A, 0x10, 0xCC)
CINZA = RGBColor(0x66, 0x66, 0x66)
PRETO = RGBColor(0x11, 0x11, 0x11)


def _base_doc(titulo: str, subtitulo: str) -> Document:
    doc = Document()
    # Fonte padrão
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = PRETO

    # Cores dos headings
    for lvl, size in (("Heading 1", 15), ("Heading 2", 12.5)):
        st = doc.styles[lvl]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = ROXO
        st.font.bold = True

    # Título
    t = doc.add_paragraph()
    r = t.add_run(titulo)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = PRETO
    sub = doc.add_paragraph()
    rs = sub.add_run(subtitulo)
    rs.italic = True
    rs.font.size = Pt(10.5)
    rs.font.color.rgb = CINZA
    doc.add_paragraph()
    return doc


def intro(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.size = Pt(10.5)
    r.font.color.rgb = CINZA


def h1(doc, t): doc.add_heading(t, level=1)
def h2(doc, t): doc.add_heading(t, level=2)


def campo(doc, label, dica=""):
    """Campo preenchível: rótulo em negrito + linha em branco com dica em cinza."""
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True
    if dica:
        p2 = doc.add_paragraph()
        rd = p2.add_run(dica)
        rd.italic = True
        rd.font.size = Pt(9.5)
        rd.font.color.rgb = CINZA
    linha = doc.add_paragraph()
    linha.add_run("__________________________________________________________________________").font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


def linha_simples(doc, texto):
    p = doc.add_paragraph()
    p.add_run(texto)


def tabela(doc, headers, linhas, extra_vazias=2):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # sombrear célula de cabeçalho
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "7A10CC")
        tcPr.append(shd)
    for linha in linhas:
        cells = t.add_row().cells
        for i, val in enumerate(linha):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(9.5)
            if not val:
                continue
    for _ in range(extra_vazias):
        cells = t.add_row().cells
        for c in cells:
            c.paragraphs[0].add_run(" ").font.size = Pt(9.5)
    doc.add_paragraph()


def rodape_uso(doc, itens):
    h1(doc, "Quem usa este template e como")
    for titulo, txt in itens:
        p = doc.add_paragraph()
        r = p.add_run(titulo + ": ")
        r.bold = True
        p.add_run(txt)


def rodape_fonte(doc):
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Aliança de IA para a Educação · Instituto Jataí — template derivado do "
                  "Guia de Contratação Pública de Soluções de IA para a Educação. "
                  "Não substitui assessoria jurídica.")
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = CINZA


# ════════════════════════════════════════════════════════════════
# RIPD SIMPLIFICADO
# ════════════════════════════════════════════════════════════════
def gerar_ripd():
    doc = _base_doc(
        "RIPD Simplificado",
        "Relatório de Impacto à Proteção de Dados para Soluções de IA na Educação",
    )
    intro(doc, "O RIPD é um instrumento previsto na LGPD (art. 38) que descreve os processos de "
               "tratamento de dados pessoais que podem gerar riscos às liberdades civis e aos direitos "
               "fundamentais. Para soluções de IA na educação é especialmente relevante porque envolve "
               "dados de crianças e adolescentes e tratamento por inferência algorítmica. Preencha os "
               "campos abaixo. Este template não substitui assessoria jurídica.")
    doc.add_paragraph()

    campo(doc, "Solução / Empresa (CNPJ) / DPO (nome, e-mail, telefone) / Data / Versão")

    h1(doc, "Parte 1 — Identificação do tratamento")
    campo(doc, "1.1 Descrição da solução", "O que faz, como usa IA, finalidade pedagógica.")
    campo(doc, "1.2 Natureza do tratamento", "Operações: coleta, armazenamento, processamento, compartilhamento, exclusão.")
    campo(doc, "1.3 Escopo", "Público-alvo (faixa etária, perfis), volume estimado de titulares, abrangência geográfica, período.")
    campo(doc, "1.4 Contexto", "Ambiente escolar, uso por menores, supervisão por professores, horário escolar/em casa.")

    h1(doc, "Parte 2 — Dados pessoais tratados")
    h2(doc, "2.1 Categorias de dados")
    tabela(doc, ["Categoria", "Dados específicos", "Finalidade", "Base legal (LGPD)"],
           [["Identificação", "Nome, série, turma, escola", "Cadastro e atribuição de perfil", "Consentimento do responsável (art. 14)"],
            ["Desempenho acadêmico", "Notas, respostas, redações", "Adaptação de conteúdo e feedback", "Consentimento do responsável (art. 14)"],
            ["Interação com a plataforma", "Tempo de uso, cliques", "Análise de engajamento", "Legítimo interesse (art. 7, IX)"],
            ["Dados inferidos pela IA", "Proficiência, risco de evasão", "Personalização e predição", "Consentimento do responsável (art. 14)"]],
           extra_vazias=2)
    campo(doc, "2.2 Dados sensíveis", "Coleta biometria, saúde, raça/etnia, religião? Justifique e descreva salvaguardas, ou declare que não coleta.")
    campo(doc, "2.3 Dados de menores", "Mecanismo de consentimento dos responsáveis (formulário digital, termo, integração com a escola).")

    h1(doc, "Parte 3 — Necessidade e proporcionalidade")
    campo(doc, "3.1 Necessidade", "Para cada categoria, por que é necessária para a finalidade pedagógica.")
    campo(doc, "3.2 Proporcionalidade", "Poderia alcançar o mesmo resultado com menos dados? Medidas de minimização.")
    campo(doc, "3.3 Qualidade dos dados", "Como a precisão é garantida. Mecanismo de correção pelo titular.")

    h1(doc, "Parte 4 — Riscos")
    tabela(doc, ["Risco", "Probabilidade (B/M/A)", "Gravidade (B/M/A)", "Mitigação"],
           [["Vazamento de dados de menores", "", "", ""],
            ["Uso para finalidade não autorizada", "", "", ""],
            ["Viés algorítmico discriminatório", "", "", ""],
            ["Decisão automatizada com impacto na trajetória escolar", "", "", ""],
            ["Reidentificação de dados pseudonimizados", "", "", ""],
            ["Acesso indevido por terceiros", "", "", ""],
            ["Retenção excessiva após encerramento", "", "", ""]],
           extra_vazias=1)
    campo(doc, "Avaliação global do risco", "Baixo / Médio / Alto, com justificativa.")

    h1(doc, "Parte 5 — Medidas de segurança")
    campo(doc, "Técnicas", "Criptografia (protocolos), controle de acesso (perfis), pseudonimização, backup, testes, monitoramento de incidentes.")
    campo(doc, "Organizacionais", "Política de privacidade (URL), treinamento de equipe, acordo de confidencialidade, plano de resposta a incidentes.")

    h1(doc, "Parte 6 — Direitos dos titulares")
    campo(doc, "6.1 Canais", "Como exercer acesso, correção, exclusão, portabilidade e revogação. Canal e prazo.")
    campo(doc, "6.2 Contestação de decisão automatizada", "Procedimento e prazo conforme art. 20 da LGPD.")

    h1(doc, "Parte 7 — Compartilhamento")
    tabela(doc, ["Terceiro", "Finalidade", "Dados", "Base legal", "Salvaguardas"],
           [["ex.: AWS", "Hospedagem", "Todos os armazenados", "Operador (art. 39)", "Contrato; servidores no Brasil"]],
           extra_vazias=2)
    campo(doc, "Transferência internacional", "Os dados vão para fora do Brasil? País, base legal e salvaguardas.")

    h1(doc, "Parte 8 — Parecer e aprovação")
    campo(doc, "Parecer do DPO", "Revisou? Nível de risco residual?")
    campo(doc, "Aprovação", "Nome, cargo, data.")
    campo(doc, "Próxima revisão", "Data (sugestão: 12 meses).")

    h1(doc, "Parte 9 — Histórico de versões")
    tabela(doc, ["Versão", "Data", "Alterações", "Responsável"], [], extra_vazias=3)

    rodape_uso(doc, [
        ("EdTech", "preencha antes de ofertar; demonstra maturidade. Atualize quando mudar dados, modelo ou terceiros."),
        ("Gestor", "use para avaliar a proteção de dados dos fornecedores. Compare RIPDs pela Parte 4 (riscos)."),
        ("Intermediário", "inclua a apresentação de RIPD como requisito em editais de credenciamento."),
    ])
    rodape_fonte(doc)

    out = OUT_DIR / "ripd-simplificado-template.docx"
    doc.save(out)
    print("  ✓", out.name)


# ════════════════════════════════════════════════════════════════
# MODEL CARD EDUCACIONAL
# ════════════════════════════════════════════════════════════════
def gerar_model_card():
    doc = _base_doc(
        "Model Card Educacional",
        "Ficha técnica padronizada para documentar uma solução de IA na educação brasileira",
    )
    intro(doc, "O Model Card descreve como um modelo de IA funciona, para que serve, quais são seus "
               "limites e como foi avaliado. Para a EdTech, é instrumento de responsabilidade e "
               "diferenciação; para o gestor, permite avaliar se a solução é adequada, segura e auditável. "
               "Nem todos os campos se aplicam a todas as soluções: preencha o pertinente e marque "
               "“Não aplicável” nos demais.")
    doc.add_paragraph()

    campo(doc, "Solução / Versão / Empresa / Contato técnico (nome e e-mail)")

    h1(doc, "1. Visão geral")
    campo(doc, "1.1 Descrição da solução", "O que faz, para quem, que problema pedagógico resolve. Linguagem acessível.")
    campo(doc, "1.2 Tipo de IA utilizada", "ML supervisionado / não supervisionado / por reforço / PLN / generativa / recomendação / preditiva / outro.")
    campo(doc, "1.3 Grau de autonomia", "Assistiva / Semiautônoma / Autônoma. Especifique por funcionalidade.")
    campo(doc, "1.4 Público-alvo", "Série/etapa, disciplina, perfil.")
    campo(doc, "1.5 Alinhamento curricular", "Relação com a BNCC: áreas, competências, habilidades.")

    h1(doc, "2. Dados")
    campo(doc, "2.1 Dados de treinamento", "Fontes, idioma, período, volume, diversidade regional e socioeconômica.")
    campo(doc, "2.2 Dados de operação", "O que a solução coleta durante o uso e para qual finalidade.")
    campo(doc, "2.3 Tratamento de dados de menores", "Consentimento, minimização, pseudonimização, retenção, exclusão.")
    campo(doc, "2.4 Localização dos dados", "País/região de armazenamento e processamento. Provedor de nuvem.")

    h1(doc, "3. Arquitetura e funcionamento")
    campo(doc, "3.1 Modelo utilizado", "Descrição acessível da arquitetura; permite compreender a lógica geral sem revelar código proprietário.")
    campo(doc, "3.2 Processo de inferência", "Como vai dos dados de entrada ao resultado. Variáveis que mais influenciam.")
    campo(doc, "3.3 Atualizações", "Frequência de retreinamento, dados utilizados, validação antes do deploy.")

    h1(doc, "4. Desempenho e avaliação")
    campo(doc, "4.1 Métricas", "Acurácia, precisão, recall, F1, kappa de Cohen ou outras pertinentes. Apresente valores.")
    campo(doc, "4.2 Avaliação por subgrupo", "Resultados por região, série, gênero, raça, escola urbana/rural, nível socioeconômico.")
    campo(doc, "4.3 Testes de viés", "Testes realizados e resultados. Se não realizados, plano para fazê-lo.")
    campo(doc, "4.4 Evidências de impacto pedagógico", "Metodologia, escala, contexto, limitações.")

    h1(doc, "5. Limitações e riscos conhecidos")
    campo(doc, "5.1 Cenários de falha", "Quando o modelo erra ou tem desempenho inferior.")
    campo(doc, "5.2 Risco de alucinação", "Para generativos: risco e medidas de mitigação.")
    campo(doc, "5.3 Risco de dependência cognitiva", "Como a solução mitiga.")
    campo(doc, "5.4 Riscos não mitigados", "Lista honesta.")

    h1(doc, "6. Supervisão humana")
    campo(doc, "6.1 Pontos de decisão humana", "Quando o professor/gestor intervém. O que pode ser anulado.")
    campo(doc, "6.2 Mecanismos de contestação", "Como alunos/famílias contestam decisões da IA.")
    campo(doc, "6.3 Formação oferecida", "Treinamento para interpretação e supervisão.")

    h1(doc, "7. Portabilidade e encerramento")
    campo(doc, "7.1 Licenciamento", "SaaS, perpétua, outro.")
    campo(doc, "7.2 Portabilidade", "Formatos de exportação, procedimento, prazo.")
    campo(doc, "7.3 Encerramento", "O que acontece com os dados ao fim do contrato. Exclusão.")
    campo(doc, "7.4 Propriedade dos dados", "Declaração de que os dados pertencem à rede, não ao fornecedor.")

    h1(doc, "8. Histórico de versões")
    tabela(doc, ["Versão", "Data", "Principais alterações"], [], extra_vazias=3)

    rodape_uso(doc, [
        ("EdTech", "preencha antes de ofertar para redes públicas. Seja honesto sobre limitações. Anexe à proposta técnica."),
        ("Gestor", "use como referência para avaliar soluções. Exija documentação equivalente no Termo de Referência."),
        ("Intermediário", "use como parte da rubrica de credenciamento. Campos essenciais: 1, 2.1-2.3, 4.1, 5, 6."),
    ])
    rodape_fonte(doc)

    out = OUT_DIR / "model-card-educacional-template.docx"
    doc.save(out)
    print("  ✓", out.name)


if __name__ == "__main__":
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    print("📄 Gerando templates .docx...")
    gerar_ripd()
    gerar_model_card()
    print("✅ Templates gerados em", OUT_DIR)
